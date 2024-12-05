import streamlit as st
import torch
from pathlib import Path
from source.utils.load_model import preprocess_image, load_all, process_layout
from concurrent.futures import ThreadPoolExecutor
import asyncio
import base64
from source.utils.convert import pdf_to_images
import PyPDF2
from source.config.schemas import ChatState
from source.model.ocr import Vision, VOCR
from source.model.correct import Correction
from io import BytesIO
import time
from docx import Document
from docx.shared import Pt
import io
import zipfile
import datetime
from utils import simulate_typing
import traceback
from source.utils.app_func import (
    phanha_spinner,
    create_message_container_html,
    login,
    set_theme,
    set_page_config,
    create_upload_indicator,
    chat_history_indicator,
    ui_indicator,
    FileHandler,
)
import os
from constant import LOGDIR
from PIL import Image
import numpy as np
import cv2
from source.utils.convert_format_bbox import sort_boxes, convert_poly_to_rectangle
from source.utils.convert_format_txt import create_text_boxes, correct_format,create_new_box
from source.utils.dang_template import PhieuDangVienParser
import pytz
from source.model.chat import LLMChat

from dotenv import load_dotenv
load_dotenv(dotenv_path="./.env")

os.environ["CUDA_LAUNCH_BLOCKING"] = "1"
    
class UnifiedChat(Vision,Correction,VOCR,PhieuDangVienParser,LLMChat):
    DEFAULT_TIMEOUT = 120
    SUPPORTED_FORMATS = ["jpg", "png", "jpeg", "pdf"]
    
    def __init__(self, models):
        self.initialize_state()
        self.models = models
        self.paddle = models['paddle']
        self.layout = models['layout']
        self.processor = models['processor']
        self.vision = models['vision']
        self.llm = models['llm']
        self.tokenizer = models['tokenizer']
        self.model_choice = None
        Vision.__init__(self, self.vision, self.processor)
        VOCR.__init__(self,self.paddle,self.model_choice)
        Correction.__init__(self)
        PhieuDangVienParser.__init__(self)
        LLMChat.__init__(self,self.llm,self.tokenizer)
        self.handler = FileHandler()
        self.executor = ThreadPoolExecutor(max_workers=3)
        self.ai_thread = None
        self.chat_container = None
        self.power_state = None
        self.history = None

    def initialize_state(self):
        if 'chat_state' not in st.session_state:
            st.session_state.chat_state = ChatState()
        self.chat_state = st.session_state.chat_state

    def reset_state(self):
        st.session_state.chat_state = ChatState()
        self.chat_state = st.session_state.chat_state
        st.rerun()
            
    def _update_chat_history(self, filename, result_text):
        self.chat_state.file_data.setdefault(filename, []).append(result_text)
        self.chat_state.correction_status[filename] = False
        content = {"role": "assistant", "content": result_text}
        self.chat_state.messages.append(content)
        if self.history:
            self.chat_state.history.append(content)
        if filename not in self.chat_state.correction_queue:
            self.chat_state.correction_queue.append(filename)
            
    async def process_corrections(self):
        try:
            filename = self.chat_state.correction_queue.pop(0)
            if filename not in self.chat_state.processed_corrections:
                file_texts = self.chat_state.file_data[filename]
                for i, text in enumerate(file_texts):
                    progress_text = f"Đang xử lý phần {i + 1}/{len(file_texts)} của {filename}"
                    st.info(progress_text)
                    try:
                        corrected_text = await self.correct_text(text)
                        if corrected_text:
                            file_texts[i] = corrected_text.text
                            message_content = self.format_correction_message(
                                corrected_text.text,
                                corrected_text.notes,
                                i + 1,
                                len(file_texts)
                            )
                            content = {"role": "assistant", "content": message_content}
                            self.chat_state.messages.append(content)
                    except Exception as e:
                        st.error(f"Lỗi khi sửa văn bản: {e}")

                self.chat_state.processed_corrections.add(filename)
                self.chat_state.correction_status[filename] = True
        except IndexError:
            st.info("Không có văn bản nào cần sửa.")
        except Exception as e:
            st.error(f"Lỗi trong quá trình sửa văn bản: {e}")
            
    def corrector(self):
        if self.chat_state.correction_in_progress:
            st.warning("Đang sửa văn bản, vui lòng đợi...")
            return

        self.chat_state.correction_in_progress = True
        with phanha_spinner("Đang sửa văn bản..."):
            asyncio.run(self.process_corrections())
        self.chat_state.correction_in_progress = False
        st.rerun()

    def _format_file_content(self, basename, texts, header):
        try:
            text = '\n'.join(texts)
            text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
            
            vietnam_timezone = pytz.timezone("Asia/Ho_Chi_Minh")
            current_time = datetime.datetime.now(vietnam_timezone).strftime('%Y-%m-%d %H:%M:%S')
            formatted_header = f"#{header}\n **{current_time}** - **{basename}**"
            
            return f"{formatted_header}\n\n{text}"
        except Exception as e:
            return f"Error formatting content for {basename}: {str(e)}"
    
    def _create_word_document(self, content):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        for line in content.split('\n'):
            if line.startswith('=' * 50):
                doc.add_heading(line.strip('='), level=1)
            else:
                doc.add_paragraph(line)
        return doc

    def create_combined_download_zip(self):
        try:
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                all_original_content = []
                all_corrected_content = []
                vietnam_timezone = pytz.timezone("Asia/Ho_Chi_Minh")
                timestamp = datetime.datetime.now(vietnam_timezone).strftime('%Y-%m-%d %H:%M:%S')
                
                for file_name, texts in self.chat_state.file_data.items():
                    basename = Path(file_name).stem
                    
                    original_content = self._format_file_content(basename, texts, "Văn bản gốc")
                    all_original_content.append(f"\n{'='*50}\n{file_name}\n{'='*50}\n")
                    all_original_content.append(original_content)
                    
                    if self.chat_state.correction_status.get(file_name, False):
                        corrected_content = self._format_file_content(basename, texts, "Văn bản đã sửa")
                        all_corrected_content.append(f"\n{'='*50}\n{file_name}\n{'='*50}\n")
                        all_corrected_content.append(corrected_content)
                
                if all_original_content:
                    combined_original = "\n".join(all_original_content)
                    txt_original_filename = f"Original_{timestamp}.txt"
                    zip_file.writestr(f"txt/{txt_original_filename}", combined_original)
                    
                    doc = self._create_word_document(combined_original)
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_original_filename = f"Original_{timestamp}.docx"
                    zip_file.writestr(f"docx/{docx_original_filename}", docx_buffer.getvalue())
                
                if all_corrected_content:
                    combined_corrected = "\n".join(all_corrected_content)
                    txt_corrected_filename = f"Corrected_{timestamp}.txt"
                    zip_file.writestr(f"txt/{txt_corrected_filename}", combined_corrected)
                    
                    doc = self._create_word_document(combined_corrected)
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_corrected_filename = f"Corrected_{timestamp}.docx"
                    zip_file.writestr(f"docx/{docx_corrected_filename}", docx_buffer.getvalue())
            
            zip_buffer.seek(0)
            return zip_buffer
            
        except Exception as e:
            st.error(f"Lỗi khi tạo file zip: {str(e)}")
            return None

    def render_download_section(self):
        st.sidebar.markdown("### Tải về văn bản")
        
        zip_buffer = self.create_combined_download_zip()
        if zip_buffer:
            st.sidebar.download_button(
                label="Tải về (TXT & DOCX)",
                data=zip_buffer,
                file_name="vanban.zip",
                mime="application/zip",
                use_container_width=True
            )          
                
    def setup_ui(self):
        ui_indicator()
        with st.sidebar:
            sidebar_logo_html = """
                <div style="margin-bottom: 20px;">
                    <svg width="200" height="70" xmlns="http://www.w3.org/2000/svg">
                        <a href="https://PhanHa.com" target="_blank">
                            <text x="10" y="50" font-size="48" font-weight="bold"
                                fill="url(#textGradient)"
                                style="font-family: Arial, sans-serif;">
                                PhanHa
                            </text>
                        </a>
                        <defs>
                            <linearGradient id="textGradient" x1="0%" y1="0%" x2="100%" y2="100%">
                                <stop offset="0%" style="stop-color:#00d4ff; stop-opacity:1" />
                                <stop offset="50%" style="stop-color:#0073e6; stop-opacity:1" />
                                <stop offset="100%" style="stop-color:#002766; stop-opacity:1" />
                            </linearGradient>
                        </defs>
                    </svg>
                </div>
            """
            st.markdown(sidebar_logo_html, unsafe_allow_html=True)
            st.markdown("---")
        if not len(self.chat_state.messages) > 1:
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                gradient_text_html = """
                    <style>
                        .gradient-text {
                            font-weight: 700;
                            font-size: 2.8em;
                            font-family: 'Arial', sans-serif;
                            text-align: center;
                            background: linear-gradient(
                                120deg,
                                #00d4ff 0%,
                                #0073e6 25%,
                                #002766 50%,
                                #0073e6 75%,
                                #00d4ff 100%
                            );
                            background-size: 200% auto;
                            color: transparent;
                            -webkit-background-clip: text;
                            background-clip: text;
                            -webkit-text-fill-color: transparent;
                            animation: shine 3s linear infinite;
                        }

                        @keyframes shine {
                            to {
                                background-position: 200% center;
                            }
                        }

                        .gradient-text:hover {
                            animation: shine 1.5s linear infinite;
                            transform: scale(1.05);
                            transition: transform 0.3s ease;
                        }
                    </style>
                    <div class="gradient-text">VISS AI</div>
                """
                st.markdown(gradient_text_html, unsafe_allow_html=True)

        self.chat_container_placeholder = st.empty()
        with self.chat_container_placeholder.container():
            self.chat_container = st.container()
            st.markdown(
                """
                <style>
                /* Chat container height and scroll */
                [data-testid="stVerticalBlock"] > [style*="flex-direction: column;"] > [data-testid="stVerticalBlock"] {
                    height: calc(100vh - 200px) !important;
                    overflow-y: auto !important;
                }
                
                /* Style the chat input container */
                .stChatInputContainer {
                    display: flex !important;
                    gap: 10px;
                    align-items: center;
                    padding-right: 40px;
                }
                
                /* Make chat input take remaining width */
                .stChatInputContainer > div:first-child {
                    flex: 1;
                }
                
                /* Style stop button container */
                button[data-testid="stop_button"] {
                    position: relative !important;
                    bottom: 0;
                    right: 0;
                    height: 40px;
                    min-width: 40px !important;
                    padding: 0 !important;
                    margin-left: 8px;
                }
                
                /* Ensure sidebar buttons maintain their original positioning */
                .st-emotion-cache-16idsys {
                    position: relative !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

    def setup_sidebar(self):
        with st.sidebar:
            system_message_editable = 'Bạn là một trợ lý AI hữu ích'
            with st.expander('⚡ Tùy chỉnh thêm'):
                temperature = st.slider('Temperature', min_value=0.01, max_value=2.0, value=0.5, step=0.05,
                                        help='Kéo sang phải để AI trả lời đa dạng hơn, kéo sang trái để trả lời chặt chẽ hơn.')
                
                top_p = st.slider('Top P', min_value=0.0, max_value=1.0, value=0.9, step=0.05,
                                help='Kéo sang trái để AI tập trung vào những câu trả lời chắc chắn nhất.')
                
                repetition_penalty = st.slider('Giảm lặp từ', min_value=1.0, max_value=1.5, value=1.1, step=0.02,
                                            help='Kéo sang phải để giảm việc AI lặp lại từ ngữ không cần thiết.')
                
                max_length = st.slider('Max Length', min_value=0, max_value=4096, value=1024, step=128,
                                    help='Điều chỉnh độ dài tối đa cho mỗi câu trả lời của AI.')
                
                top_k = st.slider('Top K', min_value=1, max_value=100, value=50, step=1,
                help='Kéo sang phải để câu trả lời đa dạng hơn, nhưng có thể ít chính xác. Kéo sang trái để câu trả lời chính xác hơn nhưng ít đa dạng.')
            with st.expander('⚙️ Thiết lập ban đầu'):
                persona_rec = st.text_area('Hướng dẫn cho AI', value=system_message_editable,
                                        help='Đây là những chỉ dẫn cơ bản để AI hiểu và trả lời đúng ý bạn ngay từ đầu cuộc trò chuyện.',
                                        height=200)
            configs = {
                "persona_rec" : persona_rec,
                "temperature": temperature,
                "top_p": top_p,
                "repetition_penalty": repetition_penalty,
                "max_new_tokens": max_length,
                "top_k": top_k
            }
            self.chat_state.configs = configs
            if st.button("Xóa lịch sử", use_container_width=True, icon="🗑️"):
                self.reset_state()
            
            if self.chat_state.file_data:
                if st.button("Tải văn bản", 
                            use_container_width=True,
                            icon="📥"):
                    self.render_download_section()
                    
                correction_button = st.button(
                    "Sửa chính tả",
                    disabled=self.chat_state.correction_in_progress,
                    use_container_width=True,
                    icon="📝"
                )
                if correction_button:
                    self.corrector()
                
                if self.chat_state.correction_status:
                    st.write("Trạng thái sửa chính tả:")
                    for filename, status in self.chat_state.correction_status.items():
                        status_text = "✅ Đã sửa" if status else "⏳ Chưa sửa"
                        st.markdown(
                        f"""
                        <div style='max-width: 250px; word-wrap: break-word; white-space: normal;'>
                            {filename}: {status_text}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
    
    def handle_user_input(self):
        # self.llm_setup_config(configs=self.chat_state.configs)
        self.vision_setup_config(configs=self.chat_state.configs)
        
        with st.sidebar:
            saving_button = st.toggle("☁️ Auto Save", value=True, key="auto_save")
            power_state = st.toggle('🛠️ Parse', value=False, key="power_toggle")
            history = st.toggle('📜 History', value=False, key="history_toggle")
            model_choice = st.selectbox(
                "📊 Chọn mô hình",
                [
                    "ChatBot",
                    "Both",
                    "VOCR_trans",
                ],
                key="select_model"
            )
            self.model_choice = model_choice
            self.power_state = power_state
            self.history = history
            uploaded_files = st.file_uploader(
                type=self.SUPPORTED_FORMATS,
                key="file_uploader",
                label='📂 Tải ảnh lên',
                accept_multiple_files=True,
            )
        if saving_button:
            self.handler.auto_save_txt(file_data=self.chat_state.file_data,
                                    correction_status=self.chat_state.correction_status,
                                    process_all=True
        )
        self.setup_sidebar()
        prompt = st.chat_input("Nhập tin nhắn hoặc tải lên ảnh...💬")
        st.button("⏹️", key="stop_button")
        
        if prompt or uploaded_files:
            self._process_user_message(prompt, uploaded_files)


    def _process_user_message(self, prompt, uploaded_files):
        with self.chat_container:
            if len(uploaded_files) == 1:
                self.display_uploaded_file(uploaded_files[0])
            elif len(uploaded_files) > 1:
                col1, col2 = st.columns([6, 1])
                with col2:
                    create_upload_indicator(uploaded_files)
            if self.model_choice == "":
                st.info("Vui lòng chọn mô hình để xử lý")
                return
            if prompt:
                self.chat_state.messages.append({
                    "role": "user",
                    "content": prompt,
                    **({"file": uploaded_files} if uploaded_files else {})
                })
                if self.history:
                    self.chat_state.history.append({
                        "role": "user",
                        "content": prompt
                    })

                st.markdown(
                    create_message_container_html(
                        prompt if prompt else "",
                        True,
                    ),
                    unsafe_allow_html=True
                )
            if len(uploaded_files) > 0 and prompt:
                if self.model_choice == "ChatBot":
                    self._process_uploaded_file(uploaded_files, prompt)
                elif self.model_choice in ['VOCR_trans', 'VOCR_seq']:
                    self.paddle_process(uploaded_files)
                elif self.model_choice == 'Both':
                    self.both_process(uploaded_files, prompt)
            elif prompt:
                if self.model_choice == "ChatBot":
                    self._handle_assistant_response(prompt)
                elif self.model_choice in ['VOCR_trans', 'VOCR_seq']:
                    st.info("Vui lòng tải lên ảnh để xử lý")
                elif self.model_choice == 'Both':
                    self.both_process(None, prompt)

    def display_uploaded_file(self, uploaded_file):
        try:
            file_name = uploaded_file.name
            file_type = file_name.split('.')[-1].lower()
            file_data = uploaded_file.getvalue()

            view_state_key = f'view_state_{file_name}'
            if view_state_key not in st.session_state:
                st.session_state[view_state_key] = 'thumbnail'

            col1, col2 = st.columns([8, 1])  

            with col2:
                if file_type == 'pdf':
                    st.markdown("""
                        <div style="background-color: #666666; color: white; padding: 8px 12px; border-radius: 15px; margin-bottom: 10px; text-align: center;">
                            <span style="font-size: 14px;">📄 PDF File</span>
                        </div>
                    """, unsafe_allow_html=True)
                else:
                    st.image(file_data, width=100)

                if file_type == 'pdf':
                    options = st.selectbox(
                        "Thao tác với PDF",
                        [
                            "",
                            "📄 Xem trước PDF",
                            "📝 Trích xuất Văn bản",
                            "ℹ️ Lấy Thông tin PDF",
                        ],
                        key=f"pdf_action_{file_name}"
                    )
                else:
                    options = st.selectbox(
                        "Thao tác với ảnh",
                        [
                            "",
                            "🔍 Xem Kích thước Đầy đủ",
                        ],
                        key=f"image_action_{file_name}"
                    )

            with col1:
                if file_type == 'pdf' and options == "📄 Xem trước PDF":
                    base64_pdf = base64.b64encode(file_data).decode('utf-8')
                    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="500" type="application/pdf"></iframe>'
                    st.markdown(pdf_display, unsafe_allow_html=True)
                
                elif file_type == 'pdf' and options == "📝 Trích xuất Văn bản":
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_data))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    
                    st.text_area("Văn bản Đã Trích xuất", text, height=300)
                    st.download_button(
                        "💾 Tải xuống Văn bản Đã Trích xuất",
                        text,
                        file_name=f"{file_name.rsplit('.', 1)[0]}.txt",
                        mime="text/plain",
                        icon="💾"
                    )
                
                elif file_type == 'pdf' and options == "ℹ️ Lấy Thông tin PDF":
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_data))
                    info = {
                        "Số trang": len(pdf_reader.pages),
                        "Kích thước tệp": f"{len(file_data) / 1024:.2f} KB",
                        "Metadata": pdf_reader.metadata if pdf_reader.metadata else "Không có metadata"
                    }
                    
                    for key, value in info.items():
                        st.markdown(f"""
                            <div style="background-color: #666666; color: white; padding: 8px 12px; border-radius: 15px; margin-bottom: 5px;">
                                <strong>{key}:</strong> {value}
                            </div>
                        """, unsafe_allow_html=True)

                elif file_type != 'pdf' and options == "🔍 Xem Kích thước Đầy đủ":
                    st.image(file_data, width=400) 

        except Exception as e:
            st.error(f"Lỗi hiển thị tệp {file_name if 'file_name' in locals() else 'unknown file'}: {str(e)}")

    def process_file(self,file):
        if file.type == "application/pdf":
            with st.spinner("Đang tải PDF..."):
                return pdf_to_images(file.read())
        else:
            with st.spinner("Đang tải ảnh..."):
                return [Image.open(file)]
            
    def process_image(self, image, filename, prompt=None, idx=None):
        try:
            display_name = f"{filename} (ảnh {idx + 1})" if idx is not None else filename
            result_text = self.vision_process(image, prompt, self.chat_container)
            if result_text:
                return result_text
        except Exception as e:
            st.error(f"Lỗi khi xử lý {display_name}: {str(e)}")
            return None

    def paddle_process(self, uploaded_files):
        """Process uploaded files and perform image processing."""
        if not uploaded_files:
            return
        type = len(uploaded_files)
        all_images = []
        file_names = []
        
        # Use each uploaded file's original name as filename.
        for uploaded_file in uploaded_files:
            try:
                # Process each uploaded file independently
                images = self.process_file(uploaded_file)
                all_images.extend(images)
                file_names.extend([uploaded_file.name] * len(images))
            except Exception as e:
                st.error(f"Lỗi đọc tệp {uploaded_file.name}: {str(e)}")

        total_images = len(all_images)
        if total_images == 0:
            st.warning("Không tìm thấy ảnh nào trong các tệp tải lên")
            return

        # Track progress with a message placeholder
        message_placeholder = self.chat_container.empty()
        generation_start_time = time.time()
        message_template = create_message_container_html(
            "Đang xử lý...", 
            is_user=False, 
            is_generating=True,
            generation_start_time=generation_start_time
        )
        message_placeholder.markdown(message_template, unsafe_allow_html=True)

        progress_bar = st.progress(0)
        processed_images = 0
        all_results = []

        # Process each image
        for idx, (image, file_name) in enumerate(zip(all_images, file_names)):
            try:
                result = self.paddle_processing(preprocess_image(image))
                
                if None not in result:
                    texts = [line[1][0] for line in result[0]]
                    boxes = [line[0] for line in result[0]]
                    boxes = [convert_poly_to_rectangle(box) for box in boxes]

                    aligned_boxes = boxes[:]
                    aligned_texts = []
                    previous_box = boxes[-1] if boxes else (0, 0, image.size[0], 0)

                    text_box = create_text_boxes(boxes, texts)
                    corrected_format = correct_format(text_box)

                    for i, text in enumerate(corrected_format):
                        if i < len(boxes):
                            aligned_texts.append((boxes[i], text))
                        else:
                            new_box = create_new_box(previous_box, image.size[1])
                            aligned_texts.append((new_box, text))
                            aligned_boxes.append(new_box)
                            previous_box = new_box

                    combined_text = []
                    for box, text in aligned_texts:
                        x1, y1, x2, y2 = box
                        combined_text.append(((x1, y1, x2, y2), text))
                    combined_text.sort(key=lambda x: x[0][1])

                    final_text = ""
                    for _, text in combined_text:
                        final_text += text
                    self._update_chat_history(file_name, final_text)  # Use original file_name
                    all_results.append(final_text)
                    processed_images += 1
                
                progress = (idx + 1) / total_images
                progress_bar.progress(progress)
                
            except Exception as e:
                st.error(f"Lỗi xử lý ảnh {idx + 1} từ tệp {file_name}: {str(e)}")
                continue

        progress_bar.empty()

        if processed_images > 0:
            combined_text = "\n".join(all_results)
            if self.power_state:
                format_output = self.format_output(combined_text)
                full_output = f"**Kết quả phân tích từ hình ảnh:**\n\n{combined_text} \n\n**Kết quả các trường:**\n{format_output}"
            else:
                full_output = combined_text
            message_template = create_message_container_html(
                full_output, 
                is_user=False, 
                is_generating=False,
                generation_start_time=generation_start_time
            )
            message_placeholder.markdown(message_template, unsafe_allow_html=True)
        else:
            message_template = create_message_container_html(
                "Không tìm thấy văn bản trong ảnh", 
                is_user=False, 
                is_generating=False,
                generation_start_time=generation_start_time
            )
            message_placeholder.markdown(message_template, unsafe_allow_html=True)
            st.error("Không có ảnh nào được xử lý thành công")

            
    def _process_uploaded_file(self, uploaded_files, prompt):
        if not uploaded_files:
            return
        name = uploaded_files[0].name
        all_images = []
        file_names = []  
        vietnam_timezone = pytz.timezone("Asia/Ho_Chi_Minh")
        current_time = datetime.datetime.now(vietnam_timezone).strftime('%Y-%m-%d %H:%M:%S')
        filename = f"{current_time} - {name} \n #Tên sẽ là tên của ảnh đầu tiên"
        for uploaded_file in uploaded_files:
            try:
                images = self.process_file(uploaded_file)
                all_images.extend(images)
                file_names.extend(uploaded_file.name)
            except Exception as e:
                st.error(f"Lỗi đọc tệp {uploaded_file.name}: {str(e)}")
        
        total_images = len(all_images)
        if total_images == 0:
            st.warning("Không tìm thấy ảnh nào trong các tệp tải lên")
            return

        progress_bar = st.progress(0)
        processed_images = 0
        all_results = []

        for idx, (image, file_name) in enumerate(zip(all_images, file_names)):

            try:
                processed_image = preprocess_image(image)
                if self.history:
                    prompt = list(self.chat_state.history)
                result = self.process_image(processed_image, file_name, prompt, idx)
                all_results.append(result)
                processed_images += 1
                progress = (idx + 1) / total_images
                self._update_chat_history(filename, result)
                self.chat_state.correction_status[filename] = False
                progress_bar.progress(progress)
                
            except Exception as e:
                st.error(f"Lỗi xử lý ảnh {idx + 1} từ tệp {file_name}: {str(e)}")
                continue
        
        progress_bar.empty()

        if processed_images > 0:
            combined_text = "\n".join(all_results)
            
            if self.power_state:
                format_output = self.format_output(combined_text)
                combined_output = f"**Kết quả phân tích từ hình ảnh:**\n\n{combined_text} \n\n**Kết quả các trường:**\n{format_output}"
            else:
                combined_output = combined_text

            message_placeholder = self.chat_container.empty()
            message_template = create_message_container_html(
                combined_output,
                is_user=False,
                is_generating=False,
            )
            message_placeholder.markdown(message_template, unsafe_allow_html=True)
        else:
            st.error("Không có ảnh nào được xử lý thành công")

    def _handle_assistant_response(self, prompt):
        with self.chat_container:
            try:
                response = self.llm_process(
                    prompt=prompt,
                    chat_container=self.chat_container 
                )
                self.chat_state.messages.append(
                    {"role": "assistant", "content": response}
                )
            except Exception as e:
                st.error(f"Lỗi tạo phản hồi: {str(e)}")

    def both_process(self, uploaded_files, prompt):
        if not hasattr(uploaded_files[0], "getvalue"):
            self._handle_assistant_response(prompt)
            return
        def process_table_area(area, prompt):
            if area is None or area.size == 0:
                return ""
                
            try:
                if isinstance(area, np.ndarray):
                    area = Image.fromarray(area)
                
                result = self.generate_content(
                    self.chat_container,
                    area,
                    prompt
                )
                return result if result else ""
            except Exception as e:
                traceback.print_exc()
                st.warning(f"Lỗi xử lý bảng: {str(e)}")
                return ""

        def process_non_table_area(area):
            if area is None or (isinstance(area, np.ndarray) and (area.size == 0 or not np.any(area))):
                return []
                
            try:
                if isinstance(area, np.ndarray):
                    area = Image.fromarray(area)
                        
                result = self.paddle_processing(area)
                if not result or not result[0]:
                    return []

                texts = [line[1][0] for line in result[0]]
                boxes = [line[0] for line in result[0]]
                
                boxes = sort_boxes(boxes)
                boxes = [convert_poly_to_rectangle(box) for box in boxes]
                aligned_boxes = boxes[:]
                aligned_texts = []
                previous_box = boxes[-1] if boxes else (0, 0, area.size[0], 0)
                text_box = create_text_boxes(boxes, texts)
                corrected_format = correct_format(text_box)
                
                for i, text in enumerate(corrected_format):
                    if i < len(boxes):
                        aligned_texts.append((boxes[i], text))
                    else:
                        new_box = create_new_box(previous_box, area.size[1])
                        aligned_texts.append((new_box, text))
                        aligned_boxes.append(new_box)
                        previous_box = new_box

                return aligned_texts
            except Exception as e:
                traceback.print_exc()
                st.warning(f"Lỗi xử lý văn bản: {str(e)}")
                return []

        vietnam_timezone = pytz.timezone("Asia/Ho_Chi_Minh")
        current_time = datetime.datetime.now(vietnam_timezone).strftime('%Y-%m-%d %H:%M:%S')
        filename = f"{current_time} - {uploaded_files[0].name}"

        all_results = []
        progress_bar = st.progress(0)
        processed_images = 0
        total_images = len(uploaded_files)

        try:
            for idx, uploaded_file in enumerate(uploaded_files):
                try:
                    # Extract images from the file
                    images = self.process_file(uploaded_file)
                    if not images:
                        st.warning(f"Tệp {uploaded_file.name} không chứa ảnh nào")
                        continue

                    for image_idx, image in enumerate(images):
                        processed_image = preprocess_image(image)
                        if processed_image is None:
                            st.error(f"Lỗi xử lý ảnh {image_idx + 1} từ tệp {uploaded_file.name}")
                            continue

                        image_array = np.array(processed_image)
                        layout = process_layout(image_array, self.layout)
                        table_areas = layout.get("Table", [])

                        final_text = ""
                        if table_areas:
                            table_text = process_table_area(image_array, prompt)
                            if table_text:
                                final_text = table_text
                        else:
                            non_table_texts = process_non_table_area(image_array)
                            final_text = "".join(text for _, text in non_table_texts)

                        if final_text:
                            all_results.append(final_text)
                            processed_images += 1

                    # Update progress
                    progress = (idx + 1) / total_images
                    progress_bar.progress(progress)

                except Exception as e:
                    st.error(f"Lỗi xử lý tệp {uploaded_file.name}: {str(e)}")
                    continue

            progress_bar.empty()

            if processed_images > 0:
                # Combine all results into a single formatted output
                combined_text = "\n".join(all_results)
                format_output = self.parse(combined_text, type=len(uploaded_files))

                formatted_dict_output = "### Định dạng đã xử lý:\n"
                for key, value in format_output.items():
                    formatted_dict_output += f"- **{key}:** \n {value if value else 'Không có dữ liệu'}\n"

                combined_output = f"### Kết quả phân tích từ hình ảnh:\n\n{combined_text}\n---\n{formatted_dict_output}"

                # Display result
                message_placeholder = self.chat_container.empty()
                message_template = create_message_container_html(
                    combined_output,
                    is_user=False,
                    is_generating=False,
                )
                message_placeholder.markdown(message_template, unsafe_allow_html=True)

                # Save result to history
                self._update_chat_history(filename, combined_output)
                self.chat_state.correction_status[filename] = False
                st.success(f"Đã xử lý thành công {processed_images}/{total_images} ảnh")
            else:
                st.error("Không có ảnh nào được xử lý thành công")

        except Exception as e:
            traceback.print_exc()
            st.error(f"Lỗi xử lý chung: {str(e)}")

    def display_chat_history(self):
        with self.chat_container:
            chat_history_indicator()
            if len(self.chat_state.messages) > 20:
                st.info("Cuộc trò chuyện quá dài, > 20 turns sẽ không hiện nữa")
            else:
                for message in self.chat_state.messages:
                    is_user = message["role"] == "user"
                    content = message.get("content", "")
                    file = message.get("file")
                    if file:
                        file_type = file[0].name.split('.')[-1].lower()
                        
                        if file_type == 'pdf':
                            col1, col2 = st.columns([8, 1])
                            with col2:
                                st.markdown(f"""
                                    <div class="message-container" style="justify-content: flex-end;">
                                        <div class="pdf-icon">📄</div>
                                    </div>
                                """, unsafe_allow_html=True)
                        
                        elif file_type in ['png', 'jpg', 'jpeg', 'gif']:
                            col1, col2 = st.columns([6, 1])
                            with col2:
                                if len(file) == 1:
                                    st.image(file[0], caption="User Image", width=100)
                                else:
                                    create_upload_indicator(file)
                    
                    has_table = '|' in content if isinstance(content, str) else False
                    if has_table:
                        st.markdown('<div style="width: 100%; overflow-x: auto;">', unsafe_allow_html=True)
                    
                    if not is_user and not hasattr(st.session_state, 'first_message_displayed'):
                        simulate_typing(content)
                        st.session_state.first_message_displayed = True
                    else:
                        message_html = create_message_container_html(content, is_user, is_generating=False)
                        st.markdown(message_html, unsafe_allow_html=True)
                    
                    if has_table:
                        st.markdown('</div>', unsafe_allow_html=True)
                
        
    def run(self):
        self.setup_ui()
        self.display_chat_history()
        self.handle_user_input()

def main():
    set_page_config()
    set_theme()
    with phanha_spinner("Đang tải mô hình, vui lòng đợi..."):
        login()
        models = load_all(vision_model_name="ha684/haocr_7B_V3.1",llm_model_name=os.getenv("LLM_MODEL"))
    chat = UnifiedChat(models)
    chat.run()

if __name__ == "__main__":
    main()
